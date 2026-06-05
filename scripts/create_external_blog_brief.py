from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "fallback-pics-external-blog-content-brief.docx"
MD_PATH = OUT_DIR / "fallback-pics-external-blog-content-brief.md"

DOMAIN = "https://fallback.pics"

COMPETITORS = [
    "placehold.co",
    "picsum.photos",
    "dummyimage.com",
    "mockimg.dev",
    "placeholdpicsum.dev",
    "placehold.jp",
    "placeholdr.dev",
]

TOPICS = [
    {
        "name": "Placeholder Image API: Complete URL Syntax Guide for Developers",
        "slug": "placeholder-image-api-url-syntax-guide",
        "description": "A practical API syntax guide showing dimensions, colors, text, avatars, skeletons, and common URL patterns developers can copy into apps and docs.",
        "goal": "Make fallback.pics the clearest reference for developers searching for a placeholder image API and drive readers to the main API landing page.",
        "keywords": "placeholder image api, image placeholder api, placeholder image url, svg placeholder image",
        "competitors": "placehold.co, dummyimage.com, placeholderimage.dev, placehold.jp",
    },
    {
        "name": "How to Fix Broken Images in HTML with onerror",
        "slug": "fix-broken-images-html-onerror",
        "description": "A plain HTML guide for replacing failed image loads with a stable fallback URL using the img onerror handler.",
        "goal": "Capture implementation intent from developers who need an immediate broken-image fix and convert them to the fallback.pics API.",
        "keywords": "img onerror fallback, broken image fallback, html image fallback, image onerror",
        "competitors": "MDN, Stack Overflow, W3Schools, blog tutorials",
    },
    {
        "name": "React Image Fallback Patterns: Missing Src, Failed Loads, and Placeholders",
        "slug": "react-image-fallback-patterns",
        "description": "A React-focused guide covering missing src values, onError handlers, reusable image components, and layout-safe placeholder URLs.",
        "goal": "Own framework-specific fallback intent and support the existing React image fallback guide with deeper examples.",
        "keywords": "react image fallback, react img onerror, react placeholder image, react broken image fallback",
        "competitors": "React blogs, Stack Overflow, DEV posts, framework snippets",
    },
    {
        "name": "Next.js Image Fallbacks Without Layout Shift",
        "slug": "nextjs-image-fallbacks-without-layout-shift",
        "description": "A Next.js implementation guide for fallback behavior, width and height handling, remote images, and avoiding CLS when images fail.",
        "goal": "Attract Next.js teams that care about production UX and Core Web Vitals, then link them into the API and guide pages.",
        "keywords": "nextjs image fallback, next image fallback, nextjs placeholder image, nextjs broken image",
        "competitors": "Vercel docs, Next.js community posts, Stack Overflow",
    },
    {
        "name": "Placeholder Image Generator vs Dummy Image Generator: What Developers Actually Need",
        "slug": "placeholder-image-generator-vs-dummy-image-generator",
        "description": "A comparison article explaining the difference between generic dummy images, production placeholders, and fallback image workflows.",
        "goal": "Clarify category language and rank for both placeholder generator and dummy image generator terms.",
        "keywords": "placeholder image generator, dummy image generator, dummy image, placeholder image",
        "competitors": "dummyimage.com, placehold.co, placeholderimage.dev, WebsitePlanet",
    },
    {
        "name": "Best Placeholder Image APIs for Developers: Feature-by-Feature Comparison",
        "slug": "best-placeholder-image-apis-for-developers",
        "description": "A comparison post covering major placeholder APIs, route syntax, formats, deterministic behavior, caching, and production suitability.",
        "goal": "Capture comparison intent and position fallback.pics as the production-safe option for deterministic fallback states.",
        "keywords": "best placeholder image api, placeholder image api, image placeholder service, placeholder api",
        "competitors": "placehold.co, picsum.photos, dummyimage.com, mockimg.dev, placeholdpicsum.dev",
    },
    {
        "name": "placehold.co Alternatives for Production Placeholder Images",
        "slug": "placehold-co-alternatives-production-placeholder-images",
        "description": "A competitor alternative article focused on when teams need deterministic fallback images instead of only quick development placeholders.",
        "goal": "Capture switching and alternative intent from placehold.co users and direct them to fallback.pics examples.",
        "keywords": "placehold.co alternative, placehold alternative, placeholder image alternative, placeholder image api",
        "competitors": "placehold.co",
    },
    {
        "name": "DummyImage Alternatives: Modern Dummy Image URLs for Web Apps",
        "slug": "dummyimage-alternatives-modern-dummy-image-urls",
        "description": "A modern comparison for developers who use dummyimage.com but need branded SVG placeholders, framework examples, or production fallbacks.",
        "goal": "Capture dummyimage alternative traffic and show practical reasons to use fallback.pics.",
        "keywords": "dummyimage alternative, dummy image generator, dummy image url, dummy images",
        "competitors": "dummyimage.com, dummy-image-generator.com, dummyimg.in",
    },
    {
        "name": "Lorem Picsum vs SVG Placeholder Images: When Random Photos Hurt UX",
        "slug": "lorem-picsum-vs-svg-placeholder-images",
        "description": "A comparison of random photo placeholders and deterministic SVG fallback images for product, docs, and dashboard interfaces.",
        "goal": "Win users who are deciding between photo placeholders and controlled fallback states.",
        "keywords": "lorem picsum alternative, picsum photos alternative, random image placeholder, svg placeholder image",
        "competitors": "picsum.photos, placeholdpicsum.dev",
    },
    {
        "name": "Product Image Placeholder Strategy for Ecommerce Catalogs",
        "slug": "product-image-placeholder-ecommerce-catalogs",
        "description": "A use-case guide for ecommerce teams handling missing supplier photos, failed CDN URLs, and incomplete product catalogs.",
        "goal": "Target high-value production use cases where stable placeholders protect catalog UX and conversion.",
        "keywords": "product image placeholder, ecommerce image placeholder, missing product image, product photo fallback",
        "competitors": "ecommerce platform docs, Shopify/WooCommerce tutorials, generic UX blogs",
    },
    {
        "name": "How to Prevent Layout Shift from Missing Images",
        "slug": "prevent-layout-shift-missing-images",
        "description": "A performance guide showing how dimensions, aspect ratios, placeholders, and fallbacks prevent content jumps when images load or fail.",
        "goal": "Connect fallback.pics to Core Web Vitals and UX quality for engineering teams.",
        "keywords": "image layout shift, prevent image layout shift, CLS images, image placeholder layout shift",
        "competitors": "web.dev, performance blogs, frontend tutorials",
    },
    {
        "name": "Skeleton Placeholder Images: When to Use Skeletons vs Static Fallbacks",
        "slug": "skeleton-placeholder-images-vs-static-fallbacks",
        "description": "An explainer that separates loading skeletons from missing-media fallbacks and shows when each pattern is appropriate.",
        "goal": "Support the skeleton generator page and help readers choose the right visual state.",
        "keywords": "skeleton placeholder generator, skeleton image placeholder, loading placeholder image, skeleton loader image",
        "competitors": "UI pattern libraries, design-system blogs, placeholder services",
    },
    {
        "name": "Avatar Placeholder Generator: Initials, Colors, and Accessibility",
        "slug": "avatar-placeholder-generator-initials-colors-accessibility",
        "description": "A guide for generating initials-based avatar placeholders that remain readable, accessible, and visually consistent.",
        "goal": "Grow the avatar placeholder cluster and attract SaaS, community, and dashboard teams.",
        "keywords": "avatar placeholder generator, initials avatar generator, user avatar placeholder, profile image fallback",
        "competitors": "ui-avatars.com, avatar services, design-system docs",
    },
    {
        "name": "SVG Placeholder Images: Why They Are Fast, Cacheable, and Scalable",
        "slug": "svg-placeholder-images-fast-cacheable-scalable",
        "description": "A technical article explaining why SVG output works well for deterministic placeholders and production fallback states.",
        "goal": "Differentiate fallback.pics from raster-first and photo-first placeholder services.",
        "keywords": "svg placeholder image, svg image placeholder, cacheable placeholder image, lightweight placeholder image",
        "competitors": "placehold.co, dummyimage.com, mockimg.dev",
    },
    {
        "name": "Cache-Control for Placeholder Images: CDN and Browser Best Practices",
        "slug": "cache-control-placeholder-images-cdn-browser",
        "description": "A technical operations guide for cache headers, immutable URLs, CDN behavior, and avoiding unbounded placeholder query strings.",
        "goal": "Position fallback.pics as a production infrastructure choice, not just a design mockup utility.",
        "keywords": "cache placeholder images, image cache control, cdn placeholder images, immutable image urls",
        "competitors": "CDN docs, Cloudflare docs, image optimization blogs",
    },
    {
        "name": "Building a Self-Hosted Placeholder Image API with Cloudflare Workers",
        "slug": "self-hosted-placeholder-image-api-cloudflare-workers",
        "description": "A developer tutorial that explains the moving parts of a worker-based SVG placeholder API and when to self-host versus use fallback.pics.",
        "goal": "Capture self-hosted intent while showing the maintenance advantage of using fallback.pics.",
        "keywords": "self hosted placeholder image api, cloudflare workers image api, svg placeholder api, build placeholder image api",
        "competitors": "Cloudflare Workers tutorials, GitHub projects, developer blogs",
    },
    {
        "name": "Placeholder Images in Storybook, Playwright, and Visual Regression Tests",
        "slug": "placeholder-images-storybook-playwright-visual-regression",
        "description": "A testing workflow article for using deterministic placeholder URLs in component libraries, test fixtures, and screenshot tests.",
        "goal": "Reach engineering teams who need stable test images and repeatable UI screenshots.",
        "keywords": "test placeholder images, storybook placeholder image, playwright image placeholder, visual regression placeholder",
        "competitors": "Storybook docs, Playwright examples, testing blogs",
    },
    {
        "name": "CSS Background Image Fallbacks: Practical Patterns and Limitations",
        "slug": "css-background-image-fallbacks",
        "description": "A guide to fallback patterns for CSS background images, where normal img onerror behavior does not apply.",
        "goal": "Capture long-tail implementation searches and explain when API placeholders should be applied at the component level.",
        "keywords": "css background image fallback, background image placeholder, css image fallback, broken background image",
        "competitors": "CSS-Tricks, MDN, Stack Overflow, frontend blogs",
    },
    {
        "name": "Responsive Placeholder Images for Cards, Banners, and Grids",
        "slug": "responsive-placeholder-images-cards-banners-grids",
        "description": "A practical guide for matching placeholder dimensions to responsive cards, banners, hero images, and dense product grids.",
        "goal": "Support generator usage across common UI layouts and reduce layout shift mistakes.",
        "keywords": "responsive placeholder image, banner placeholder image, card image placeholder, image placeholder sizes",
        "competitors": "placeholder generators, UI layout tutorials, design-system blogs",
    },
    {
        "name": "OG Image Placeholders for Blogs, Docs, and Social Sharing",
        "slug": "og-image-placeholders-blogs-docs-social-sharing",
        "description": "A content workflow guide for using generated placeholder images as Open Graph and social preview assets while final artwork is unavailable.",
        "goal": "Capture adjacent content and documentation teams while demonstrating the generated image API.",
        "keywords": "og image placeholder, open graph image placeholder, social image placeholder, blog image placeholder",
        "competitors": "OG image generators, social preview tools, content ops blogs",
    },
    {
        "name": "Placeholder Images for CMS Previews and Missing Media Fields",
        "slug": "placeholder-images-cms-previews-missing-media",
        "description": "A CMS-focused article for handling empty media fields, preview states, and editorial workflows without broken image icons.",
        "goal": "Target teams with CMS, docs, and publishing workflows that need reliable fallback media.",
        "keywords": "cms image placeholder, missing media placeholder, cms preview image, image fallback cms",
        "competitors": "CMS docs, headless CMS blogs, ecommerce tutorials",
    },
    {
        "name": "Mobile App Image Fallbacks: Avatars, Cards, and Offline States",
        "slug": "mobile-app-image-fallbacks-avatars-cards-offline",
        "description": "A product UX article on fallback image states for mobile interfaces, including profile photos, cards, and unreliable networks.",
        "goal": "Extend fallback.pics beyond web-only examples while keeping the API URL workflow relevant.",
        "keywords": "mobile image fallback, app image placeholder, avatar fallback mobile, offline image fallback",
        "competitors": "mobile UX blogs, app development tutorials, design-system docs",
    },
    {
        "name": "Privacy-Safe Placeholder Images: Why URL Text and Uploads Matter",
        "slug": "privacy-safe-placeholder-images-url-text-uploads",
        "description": "A trust-focused article explaining what not to put in placeholder URLs and why no-upload deterministic placeholders can be safer.",
        "goal": "Differentiate fallback.pics on privacy, safety, and production-readiness.",
        "keywords": "privacy placeholder image api, safe placeholder image, no upload placeholder image, placeholder url privacy",
        "competitors": "AI image placeholder tools, upload-based image generators, privacy blogs",
    },
    {
        "name": "Branded Fallback Images for SaaS Dashboards and Internal Tools",
        "slug": "branded-fallback-images-saas-dashboards-internal-tools",
        "description": "A SaaS use-case guide for consistent preview, report, avatar, workspace, and empty-media states in operational interfaces.",
        "goal": "Tie fallback.pics to high-value B2B product surfaces and future paid/team features.",
        "keywords": "fallback image service, branded placeholder image, dashboard image placeholder, saas image fallback",
        "competitors": "design-system blogs, SaaS UX articles, internal tool platforms",
    },
    {
        "name": "From Broken Image Icon to Branded Fallback: A Production Rollout Checklist",
        "slug": "broken-image-icon-to-branded-fallback-checklist",
        "description": "A checklist for auditing image surfaces, choosing fallback dimensions, updating components, QA testing, and monitoring broken images.",
        "goal": "Create a practical conversion-focused asset that external writers and internal teams can use as the rollout model.",
        "keywords": "broken image fallback, broken image icon, fallback image checklist, image fallback rollout",
        "competitors": "frontend best-practice blogs, QA checklists, UX reliability articles",
    },
]


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_in: float = 6.5) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)

    for label, detail in rows:
        row = table.add_row()
        label_cell, detail_cell = row.cells
        set_cell_width(label_cell, 1.18)
        set_cell_width(detail_cell, 5.32)
        set_cell_shading(label_cell, "E8EEF5")
        set_cell_text(label_cell, label, bold=True, color="1F4D78")
        set_cell_text(detail_cell, detail)

    document.add_paragraph()


def add_topic(document: Document, index: int, topic: dict[str, str]) -> None:
    heading = document.add_heading(f"{index}. {topic['name']}", level=2)
    heading.keep_with_next = True
    url = f"{DOMAIN}/blog/{topic['slug']}/"
    add_table(
        document,
        [
            ("Blog name", topic["name"]),
            ("Short description", topic["description"]),
            ("What we want to achieve", topic["goal"]),
            ("Keywords", topic["keywords"]),
            ("Our domain", DOMAIN),
            ("Planned live URL", url),
            ("Competitor context", topic["competitors"]),
        ],
    )


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def build_docx() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("fallback.pics External Blog Content Brief")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = document.add_paragraph()
    subtitle.add_run(
        "Prepared for external content partners. Covers 25 planned blog articles, target keywords, business goals, planned live URLs, domain details, and competitor context."
    )

    add_table(
        document,
        [
            ("Domain", DOMAIN),
            ("Publishing location", f"{DOMAIN}/blog/[article-slug]/"),
            ("Brand positioning", "Developer-focused placeholder image API and fallback image service."),
            ("Primary objective", "Build topical authority around placeholder images, fallback images, dummy image generators, and implementation guides while driving qualified traffic to fallback.pics product pages."),
            ("Competitor set", ", ".join(COMPETITORS)),
        ],
    )

    document.add_heading("Writing Goals", level=1)
    add_bullets(
        document,
        [
            "Create practical, developer-friendly blog posts with copy-paste examples and clear implementation guidance.",
            "Use each post to support one primary fallback.pics landing page or implementation guide.",
            "Differentiate fallback.pics around deterministic SVG placeholders, production fallback states, cacheability, and privacy-safe no-upload workflows.",
            "Avoid unsupported claims about exact traffic or volume unless current keyword data is supplied separately.",
        ],
    )

    document.add_heading("General Requirements for Every Article", level=1)
    add_bullets(
        document,
        [
            "Mention fallback.pics naturally in the introduction and conclusion.",
            "Include the planned live URL as the canonical publishing destination.",
            "Include the primary keyword in the title, H1, introduction, meta description, and one subheading where natural.",
            "Include at least one working fallback.pics API example using the /api/v1/ route.",
            "Link internally to the assigned fallback.pics product or guide page.",
            "Do not add FAQPage or HowTo schema recommendations for these posts.",
        ],
    )

    document.add_heading("Competitor Context", level=1)
    add_table(
        document,
        [
            ("placehold.co", "Simple placeholder URLs. Compete with production fallback reliability and migration/alternative content."),
            ("picsum.photos", "Random photo placeholders. Differentiate with deterministic branded SVG fallback states."),
            ("dummyimage.com", "Classic dummy image URL syntax. Differentiate with modern framework examples and production UX guidance."),
            ("mockimg.dev", "Feature-rich generator. Avoid feature-count competition; focus on reliability and broken-image workflows."),
            ("placeholdpicsum.dev", "Broad all-in-one generator and photo API. Focus on privacy-safe, no-upload, deterministic fallback use cases."),
            ("placehold.jp", "Utility preset generator. Compete on English developer guides and production workflows."),
            ("placeholdr.dev", "AI prompt-based placeholder API. Differentiate with fast non-AI deterministic SVG generation."),
        ],
    )

    document.add_heading("25 Blog Briefs", level=1)
    for index, topic in enumerate(TOPICS, start=1):
        add_topic(document, index, topic)

    document.add_heading("External Partner Note", level=1)
    document.add_paragraph(
        "This brief is a planning document. Final article drafts should verify live URLs, product examples, and competitor references before publication."
    )

    document.save(DOCX_PATH)


def build_markdown() -> None:
    lines = [
        "# fallback.pics External Blog Content Brief",
        "",
        "Prepared for external content partners.",
        "",
        f"- Domain: {DOMAIN}",
        f"- Publishing location: {DOMAIN}/blog/[article-slug]/",
        "- Primary objective: Build topical authority around placeholder images, fallback images, dummy image generators, and implementation guides while driving qualified traffic to fallback.pics product pages.",
        f"- Competitors: {', '.join(COMPETITORS)}",
        "",
        "## 25 Blog Briefs",
        "",
    ]

    for index, topic in enumerate(TOPICS, start=1):
        url = f"{DOMAIN}/blog/{topic['slug']}/"
        lines.extend(
            [
                f"### {index}. {topic['name']}",
                "",
                f"- Blog name: {topic['name']}",
                f"- Short description: {topic['description']}",
                f"- What we want to achieve: {topic['goal']}",
                f"- Keywords: {topic['keywords']}",
                f"- Our domain: {DOMAIN}",
                f"- Planned live URL: {url}",
                f"- Competitor context: {topic['competitors']}",
                "",
            ]
        )

    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
